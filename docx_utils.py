"""
docx_utils.py — low-level python-docx helpers (borders, shading, fields).

python-docx has no public API for table/page borders, cell shading, or field
codes (PAGE / NUMPAGES), so we reach into the underlying OOXML. Kept in one
place so proposal.py stays readable.
"""
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1B, 0x25, 0x4B)
BLACK = RGBColor(0, 0, 0)


def _border_el(edge, sz="6", val="single", color="000000", space="0"):
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    return el


def set_cell_borders(cell, edges=("top", "left", "bottom", "right"),
                     sz="6", color="000000", val="single"):
    """Draw borders on the named edges of a single cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in edges:
        existing = tcBorders.find(qn(f"w:{edge}"))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(_border_el(edge, sz=sz, color=color, val=val))


def clear_table_borders(table):
    """Remove all borders from a table (start from a clean slate)."""
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tblBorders.append(_border_el(edge, val="none", sz="0"))
    tblPr.append(tblBorders)


def shade_cell(cell, fill="D9D9D9"):
    """Fill a cell with a solid background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_page_border(section, sz="18", color="000000", space="20"):
    """Box border around the whole page (offset from the page edge)."""
    sectPr = section._sectPr
    existing = sectPr.find(qn("w:pgBorders"))
    if existing is not None:
        sectPr.remove(existing)
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        pgBorders.append(_border_el(edge, sz=sz, color=color, space=space))
    sectPr.append(pgBorders)


def add_page_line(container, x_pt, y1_pt, y2_pt, weight="1pt", color="black"):
    """Draw a straight line at absolute page coordinates (points). Placed in a
    section header it repeats on every page — used for the full-height AMOUNT
    column divider that runs from the band down to the TOTAL box."""
    from docx.oxml import parse_xml
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml">'
        f'<v:line from="{x_pt}pt,{y1_pt}pt" to="{x_pt}pt,{y2_pt}pt" '
        f'strokecolor="{color}" strokeweight="{weight}" '
        'style="position:absolute;mso-position-horizontal-relative:page;'
        'mso-position-vertical-relative:page;z-index:-1"/>'
        '</w:pict>'
    )
    run._r.append(parse_xml(xml))
    return p


def add_field(paragraph, instr, default="1"):
    """Insert a Word field (e.g. 'PAGE', 'NUMPAGES') that updates on open."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {instr} "
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = default
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (begin, instrText, sep, t, end):
        run._r.append(el)
    return run


def repeat_row_as_header(row):
    """Mark a table row to repeat at the top of every page it spans."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def cant_split_row(row):
    """Keep a table row from being split across a page break."""
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement("w:cantSplit")
    cs.set(qn("w:val"), "true")
    trPr.append(cs)


def no_space(paragraph, before=0, after=0, line=None):
    """Tighten paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    return paragraph


def set_col_widths(table, widths):
    """Force column widths (python-docx ignores cell.width without this)."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = w
