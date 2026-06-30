"""
proposal.py — render a normalized proposal dict into a branded Metro .docx.

This is pure layout. build.py does the code lookups and hands us a `doc` dict
with everything already resolved to display text:

    doc = {
      "header": {for, address, phone, email, date, terms,
                 job_address (multiline), attention, bid_number, submitted_by,
                 ccb, cc, job_footer},
      "tariff_notes": [str, ...],          # italic * bullets under the band
      "intro": str,                        # "Metro Access will provide ..."
      "gate_summary": [str, ...],          # centered Install: list
      "gates": [ {title, lines:[{qty:int|None, text}]} ],
      "options": [ ... see render_options ... ],
      "notes": [str], "warranties": [str], "exclusions": [str],
      "total": str|None,
    }

The top info box and the footer (signature / address / fine print) live in the
Word page header & footer so they repeat on every page automatically. The
"WE PROPOSE ... | AMOUNT" band repeats via tblHeader. Body content flows
between, with page breaks at section boundaries.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION

import docx_utils as U

HERE = os.path.dirname(os.path.abspath(__file__))
LOGO = os.path.join(HERE, "assets", "metro-logo.png")

SERIF = "Times New Roman"
CONTENT_W = Inches(6.6)      # left content column
AMOUNT_W = Inches(1.0)       # right AMOUNT gutter
PAGE_BODY_W = CONTENT_W + AMOUNT_W
TEXT_GAP = Inches(0.28)      # keep body text clear of the AMOUNT divider line
RIGHT_INDENT = AMOUNT_W + TEXT_GAP   # right indent for full-width body paragraphs


# ----------------------------------------------------------------------------
# small paragraph/run helpers
# ----------------------------------------------------------------------------
def _run(p, text, *, bold=False, italic=False, size=10, font=SERIF,
         color=None, underline=False, allcaps=False, smallcaps=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.name = font
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    if allcaps:
        r.font.all_caps = True
    if smallcaps:
        r.font.small_caps = True
    return r


def _para(container, *, align=None, before=0, after=2, line=1.0, reserve_amount=True):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    if reserve_amount:                 # keep body text left of the AMOUNT column
        p.paragraph_format.right_indent = RIGHT_INDENT
    U.no_space(p, before=before, after=after, line=line)
    return p


# ----------------------------------------------------------------------------
# page header — the info box (repeats on every page)
# ----------------------------------------------------------------------------
def _build_header(section, h):
    header = section.header
    header.is_linked_to_previous = False
    # clear the default empty paragraph
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    # full-height AMOUNT-column divider: a page-anchored line from the band
    # (top of the body content area) down into the TOTAL box, drawn here so it
    # repeats on every page. top_margin is sized to always exceed the header
    # height, so the band sits at top_margin and the line starts right at it.
    x = section.left_margin.pt + CONTENT_W.pt
    y_top = section.top_margin.pt + 1
    y_bottom = section.page_height.pt - section.bottom_margin.pt + 17  # into TOTAL box
    U.add_page_line(header, round(x, 1), round(y_top, 1), round(y_bottom, 1))

    # --- title row: PROPOSAL | Pg X of Y | logo ---
    trow = header.add_table(rows=1, cols=3, width=PAGE_BODY_W)
    trow.alignment = WD_TABLE_ALIGNMENT.CENTER
    U.clear_table_borders(trow)
    U.set_col_widths(trow, [Inches(3.5), Inches(2.1), Inches(2.0)])

    c0 = trow.cell(0, 0).paragraphs[0]
    U.no_space(c0, after=0)
    _run(c0, "PROPOSAL", bold=True, size=30, font=SERIF)

    # boxed "Pg. X of Y" — nested 1-cell table so the border hugs the text
    midcell = trow.cell(0, 1)
    midcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    U.no_space(midcell.paragraphs[0], after=0)
    pgbox = midcell.add_table(rows=1, cols=1)
    pgbox.alignment = WD_TABLE_ALIGNMENT.CENTER
    U.set_cell_borders(pgbox.cell(0, 0), sz="6")
    c1 = pgbox.cell(0, 0).paragraphs[0]
    c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    U.no_space(c1, before=1, after=1)
    _run(c1, "Pg. ", size=11)
    U.add_field(c1, "PAGE", "1")
    _run(c1, " of ", size=11)
    U.add_field(c1, "NUMPAGES", "1")

    c2 = trow.cell(0, 2).paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    U.no_space(c2, after=0)
    if os.path.isfile(LOGO):
        c2.add_run().add_picture(LOGO, width=Inches(1.15))

    # --- fields grid: 2 columns ---
    grid = header.add_table(rows=4, cols=2, width=PAGE_BODY_W)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    U.clear_table_borders(grid)
    U.set_col_widths(grid, [Inches(4.6), Inches(3.0)])

    def field(cell, label, value, value_bold=False, hang=None):
        p = cell.paragraphs[0]
        U.no_space(p, before=3, after=3)
        if hang is not None:                  # hanging indent: wrapped/2nd lines align
            p.paragraph_format.left_indent = hang
            p.paragraph_format.first_line_indent = -hang
        _run(p, f"{label} ", bold=True, size=10)
        if value:
            for i, ln in enumerate(str(value).split("\n")):
                if i:
                    p.add_run().add_break()
                _run(p, ln, bold=value_bold, size=10)
        return p

    field(grid.cell(0, 0), "For:", h.get("for"))
    field(grid.cell(0, 1), "Phone:", h.get("phone"))
    field(grid.cell(1, 0), "Address:", h.get("address"))
    field(grid.cell(1, 1), "", "")
    # Date + Terms share a line
    p = grid.cell(2, 0).paragraphs[0]
    U.no_space(p, before=3, after=3)
    _run(p, "Date: ", bold=True, size=10)
    _run(p, f"{h.get('date','')}      ", size=10)
    _run(p, "Terms: ", bold=True, size=10)
    _run(p, h.get("terms", "") or "", size=10)
    field(grid.cell(2, 1), "Email:", h.get("email"))
    field(grid.cell(3, 0), "Job Address:", h.get("job_address"),
          value_bold=True, hang=Inches(1.0))
    field(grid.cell(3, 1), "Attention:", h.get("attention"))

    # bottom rule under the info box
    rule = header.add_paragraph()
    U.no_space(rule, before=2, after=0)
    pPr = rule._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)


# ----------------------------------------------------------------------------
# page footer — fine print + signature/address block (repeats every page)
# ----------------------------------------------------------------------------
FINE_PRINT = (
    "The above proposal is valid for 30 days.  Agreements are contingent upon "
    "strikes, accidents or other conditions beyond our control.  We carry "
    "manufacturers’, contractors’, & employers’ liability & workman’s compensation "
    "insurance. Customer agrees that all equipment is the property of contractor "
    "& allows contractor access to property to remove equipment if full payment is "
    "not made per contract terms. A 1 ½% finance charge per month is charged on all "
    "past due accounts, plus all attorney fees & court cost for collection."
)


def _build_footer(section, h, total=None):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    # --- DATE INSTALLATION DESIRED | TOTAL band ---
    band = footer.add_table(rows=1, cols=2, width=PAGE_BODY_W)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    U.clear_table_borders(band)
    U.set_col_widths(band, [CONTENT_W, AMOUNT_W])
    for c in band.rows[0].cells:
        U.set_cell_borders(c, edges=("top", "bottom"), sz="8")
    lp = band.cell(0, 0).paragraphs[0]
    U.no_space(lp, before=1, after=1)
    _run(lp, "DATE INSTALLATION DESIRED", size=10)
    rp = band.cell(0, 1).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    U.no_space(rp, before=1, after=1)
    _run(rp, "TOTAL", bold=True, size=10, smallcaps=True)
    if total:
        tp = band.cell(0, 1).add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        U.no_space(tp, before=0, after=1)
        amt = str(total)
        _run(tp, amt if amt.startswith("$") else f"${amt}", bold=True, size=11)
    U.set_cell_borders(band.cell(0, 1), edges=("top", "bottom", "left"), sz="8")

    # --- fine print + Metro logo in the bottom-right corner (like the original) ---
    fp = footer.add_paragraph()
    U.no_space(fp, before=2, after=2, line=0.9)
    _run(fp, FINE_PRINT, size=6.5)
    ids = footer.add_paragraph()
    U.no_space(ids, before=0, after=0)
    _run(ids, f"CCB # {h.get('ccb','')}   CC # {h.get('cc','')}        "
              f"{h.get('job_footer','')}", size=7)
    lg = footer.add_paragraph()
    lg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    U.no_space(lg, before=2, after=0)
    if os.path.isfile(LOGO):
        lg.add_run().add_picture(LOGO, width=Inches(1.1))

    # --- signature / address block ---
    sig = footer.add_table(rows=2, cols=2, width=PAGE_BODY_W)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    U.clear_table_borders(sig)
    U.set_col_widths(sig, [Inches(3.8), Inches(3.8)])

    a = sig.cell(0, 0).paragraphs[0]
    U.no_space(a, after=1)
    _run(a, "WE ACCEPT THE ABOVE PROPOSAL:", size=9)
    addr = sig.cell(0, 1).paragraphs[0]
    addr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    U.no_space(addr, after=0)
    _run(addr, "2525 NE COLUMBIA BLVD PORTLAND OR. 97211", bold=True, size=8)
    addr2 = sig.cell(1, 1).paragraphs[0]
    addr2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    U.no_space(addr2, before=0, after=0)
    _run(addr2, "(503) 595-4716    (503) 285-1793 Fax", bold=True, size=8)

    by = sig.cell(1, 0).paragraphs[0]
    U.no_space(by, before=4, after=0)
    _run(by, "BY: __________________________  DATE: __________", size=9)
    sub = sig.cell(1, 0).add_paragraph()
    U.no_space(sub, before=4, after=0)
    _run(sub, "Submitted By:   ", size=9)
    _run(sub, (h.get("submitted_by") or "").upper(), size=9)
    if h.get("bid_number"):
        bn = sig.cell(1, 0).add_paragraph()
        U.no_space(bn, before=0, after=0)
        _run(bn, f"                {h.get('bid_number')}", size=8)


# ----------------------------------------------------------------------------
# body content
# ----------------------------------------------------------------------------
def _band(body):
    """The 'WE PROPOSE TO FURNISH THE FOLLOWING | AMOUNT' header band."""
    t = body.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    U.clear_table_borders(t)
    U.set_col_widths(t, [CONTENT_W, AMOUNT_W])
    for c in t.rows[0].cells:
        U.set_cell_borders(c, edges=("top", "bottom"), sz="10")
    U.set_cell_borders(t.cell(0, 1), edges=("top", "bottom", "left"), sz="10")
    lp = t.cell(0, 0).paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    U.no_space(lp, before=1, after=1)
    _run(lp, "WE PROPOSE TO FURNISH THE FOLLOWING", bold=True, size=11, smallcaps=True)
    rp = t.cell(0, 1).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    U.no_space(rp, before=1, after=1)
    _run(rp, "AMOUNT", bold=True, size=10, smallcaps=True)
    U.repeat_row_as_header(t.rows[0])
    return t


QTY_TAB, DESC_TAB = Inches(0.62), Inches(1.25)


def _format_amount_text(val, deduct=False):
    """Render a value for the AMOUNT column: a percent stays as-is, a number gets
    a '$' prefix, anything else (a note like 'by others') renders verbatim."""
    import re
    s = str(val).strip()
    if not s:
        return ""
    if "%" in s:
        body = s
    elif re.fullmatch(r"\$?[\d,]+(\.\d+)?", s):
        body = f"$ {s.lstrip('$').strip()}"
    else:
        body = s
    return f"<{body}>" if deduct else body


def _format_scope_para(p, qty, text, label="", reserve_amount=True):
    """Format an existing paragraph as a 3-column scope line (label / qty / desc)."""
    U.no_space(p, before=0, after=1, line=1.0)
    pf = p.paragraph_format
    pf.left_indent = DESC_TAB
    pf.first_line_indent = -DESC_TAB
    # keep text clear of the AMOUNT divider (full body indent, or a small gap
    # when already inside a content-width cell, e.g. an option block)
    pf.right_indent = RIGHT_INDENT if reserve_amount else TEXT_GAP
    pf.tab_stops.add_tab_stop(QTY_TAB)
    pf.tab_stops.add_tab_stop(DESC_TAB)
    marker = f"{qty})" if qty not in (None, 0, "") else "—"
    _run(p, (label or "") + f"\t{marker}\t", size=10)
    _run(p, text, size=10)
    return p


def _scope_line(container, qty, text, label="", reserve_amount=True):
    """One scope line, in three aligned columns like the original:
        Install:   1)        <description>
                   1)        <description>
        Other:     —         <no-count description>
    `label` (e.g. "Install:" / "Other:") prints on the rows it's given; the qty
    and description columns stay aligned via fixed tab stops + a hanging indent
    so wrapped text lines up too. Used for gate scope and inside option blocks.
    """
    return _format_scope_para(container.add_paragraph(), qty, text,
                              label=label, reserve_amount=reserve_amount)


def _scope_amount_row(body, qty, text, label, amount, deduct=False):
    """A scope line with a per-item price/note shown in the AMOUNT column on the
    same row (e.g. 'New Enclosure ....... $450' or '... by others')."""
    t = body.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    U.clear_table_borders(t)
    U.set_col_widths(t, [CONTENT_W, AMOUNT_W])
    _format_scope_para(t.cell(0, 0).paragraphs[0], qty, text,
                       label=label, reserve_amount=False)
    rc = t.cell(0, 1)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    rp = rc.paragraphs[0]
    U.no_space(rp, before=0, after=1)
    _run(rp, _format_amount_text(amount, deduct), size=10)
    return t


def _priced_note_row(body, label, amount, deduct=False):
    """A custom priced note: right-aligned label in the content column and a
    price in the AMOUNT column on the same row (e.g. 'Not-To-Exceed Total: $540')."""
    t = body.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    U.clear_table_borders(t)
    U.set_col_widths(t, [CONTENT_W, AMOUNT_W])
    lp = t.cell(0, 0).paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # confine the label to the right ~half so a long label wraps to a narrow,
    # right-aligned column (instead of stretching across to the divider)
    lp.paragraph_format.left_indent = Inches(3.0)
    lp.paragraph_format.right_indent = TEXT_GAP   # stay clear of the divider line
    U.no_space(lp, before=2, after=2)
    if label:
        _run(lp, label, bold=True, underline=True, size=10)
    rc = t.cell(0, 1)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM   # align price to the note's last line
    rp = rc.paragraphs[0]
    U.no_space(rp, before=2, after=2)
    if amount not in (None, ""):
        _run(rp, _format_amount_text(amount, deduct), size=10)
    return t


def _page_break(body):
    p = body.add_paragraph()
    U.no_space(p, after=0)
    p.add_run().add_break(WD_BREAK.PAGE)


def render_body(body, doc):
    _band(body)

    # tariff notes (italic, * bullet)
    for n in doc.get("tariff_notes", []):
        p = _para(body, before=2, after=2)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        _run(p, "✱  ", size=9)
        _run(p, n, italic=True, bold=True, size=9.5)

    # intro sentence (bold, centered)
    if doc.get("intro"):
        p = _para(body, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2)
        _run(p, doc["intro"], bold=True, size=12)

    # Install: gate summary (centered, underlined header)
    if doc.get("gate_summary"):
        p = _para(body, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=1)
        _run(p, "Install:", bold=True, underline=True, size=11)
        for g in doc["gate_summary"]:
            gp = _para(body, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
            _run(gp, g, bold=True, size=11)

    # per-gate scope
    for gi, gate in enumerate(doc.get("gates", [])):
        title = _para(body, before=10, after=2)
        _run(title, gate["title"], bold=True, underline=True, size=10.5)
        # keep the whole gate block on one page (gives one-gate-per-page when a
        # gate is too long to fit in the remaining space, like the sample)
        block = [title]
        first_scope = True
        for ln in gate["lines"]:
            if ln.get("amount_note") is not None:
                # custom priced note: right-aligned label + price in AMOUNT column
                _priced_note_row(body, ln.get("amount_note", ""),
                                 ln.get("amount"), ln.get("deduct"))
                continue
            # only label the first counted line "Install:" (not a no-count line)
            label = "Install:" if (first_scope and
                                   ln.get("qty") not in (None, 0, "")) else ""
            first_scope = False
            if ln.get("amount") not in (None, ""):
                # scope line with a per-item price/note in the AMOUNT column
                _scope_amount_row(body, ln.get("qty"), ln["text"], label,
                                  ln["amount"], ln.get("deduct"))
                continue
            block.append(_scope_line(body, ln.get("qty"), ln["text"], label=label))
        for para in block[:-1]:
            para.paragraph_format.keep_with_next = True
        for para in block:
            para.paragraph_format.keep_together = True

    # options — flow right after the gates; only spill to a new page if there
    # isn't room (no forced page break)
    if doc.get("options"):
        render_options(body, doc["options"])

    # notes / warranties / exclusions — likewise flow after options and only
    # break to a new page when out of space
    if any(doc.get(k) for k in ("notes", "warranties", "exclusions")):
        render_nwe(body, doc)


def _detail_option(body, opt):
    """A titled option block: bold-underlined title, coded/free-text scope lines
    (Install: / qty / description), and one amount aligned in the AMOUNT column."""
    t = body.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    U.clear_table_borders(t)
    U.set_col_widths(t, [CONTENT_W, AMOUNT_W])
    left = t.cell(0, 0)
    tp = left.paragraphs[0]
    U.no_space(tp, before=2, after=2)
    if opt.get("title"):
        _run(tp, opt["title"], bold=True, underline=True, size=10.5)
    for ln in opt.get("lines", []):
        if ln.get("amount_note") is not None:
            continue   # options carry a single block amount, not per-line notes
        p = _scope_line(left, ln.get("qty"), ln["text"],
                        label=ln.get("label", ""), reserve_amount=False)
        if ln.get("bold"):
            for r in p.runs:
                r.bold = True

    rc = t.cell(0, 1)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    ap = rc.paragraphs[0]
    ap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    U.no_space(ap, after=0)
    if opt.get("amount") not in (None, ""):
        _run(ap, _format_amount_text(opt["amount"], opt.get("deduct")), size=10)
    _para(body, after=8)   # spacer between options


def render_options(body, options):
    head = _para(body, align=WD_ALIGN_PARAGRAPH.CENTER, before=14, after=6)
    head.paragraph_format.keep_with_next = True
    _run(head, "OPTIONS – Circle options chosen to be added to totals:",
         bold=True, underline=True, size=10.5)

    def amount_para(cell, amount, deduct=False):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        U.no_space(p, after=0)
        if amount:
            txt = f"<${amount}>" if deduct else f"${amount}"
            _run(p, txt, size=10)

    for opt in options:
        # detailed option: a titled block of coded/free-text scope lines + an
        # amount (e.g. "Residential Elevator:" with a Brivo panel, a reader, ...)
        if isinstance(opt, dict) and opt.get("lines") is not None and "kind" not in opt:
            _detail_option(body, opt)
            continue

        # each option is a 2-col table row: text | amount
        t = body.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        U.clear_table_borders(t)
        U.set_col_widths(t, [CONTENT_W, AMOUNT_W])
        U.set_cell_borders(t.cell(0, 1), edges=("left",), sz="6")
        left = t.cell(0, 0)
        lp = left.paragraphs[0]
        U.no_space(lp, after=2)

        if opt.get("kind") == "block":
            _run(lp, opt.get("title", ""), bold=True, underline=True, size=10)
            for b in opt.get("bullets", []):
                bp = left.add_paragraph()
                U.no_space(bp, after=0)
                bp.paragraph_format.left_indent = Inches(0.3)
                bp.paragraph_format.first_line_indent = Inches(-0.2)
                _run(bp, "—  ", size=10)
                _run(bp, b, size=10)
            # priced sub-lines: each label gets its own row so the amount on the
            # right lines up with it
            for sub in opt.get("priced", []):
                st = body.add_table(rows=1, cols=2)
                st.alignment = WD_TABLE_ALIGNMENT.CENTER
                U.clear_table_borders(st)
                U.set_col_widths(st, [CONTENT_W, AMOUNT_W])
                U.set_cell_borders(st.cell(0, 1), edges=("left",), sz="6")
                sp = st.cell(0, 0).paragraphs[0]
                U.no_space(sp, after=0)
                sp.paragraph_format.left_indent = Inches(0.6)
                _run(sp, "✱  ", size=10)
                _run(sp, sub["label"], bold=True, size=10)
                ap = st.cell(0, 1).paragraphs[0]
                ap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                U.no_space(ap, after=0)
                _run(ap, f"${sub['amount']}", size=10)
        else:
            _run(lp, "•  ", size=10)
            _run(lp, opt.get("text", ""), size=10)
            if opt.get("cost_per_gate"):
                cg = left.add_paragraph()
                cg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                U.no_space(cg, after=0)
                _run(cg, "Cost Per Gate", bold=True, size=9)
            amount_para(t.cell(0, 1), opt.get("amount"), opt.get("deduct"))
        # spacer
        sp = _para(body, after=4)


def render_nwe(body, doc):
    def section(label, items, bullet="✱"):
        if not items:
            return
        h = _para(body, before=10, after=2)
        h.paragraph_format.keep_with_next = True   # don't orphan the header
        _run(h, label, bold=True, underline=True, size=10.5)
        for it in items:
            p = body.add_paragraph()
            U.no_space(p, after=1, line=1.0)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.right_indent = RIGHT_INDENT
            _run(p, f"{bullet}  ", size=9.5)
            _run(p, it, size=9.5)

    section("Notes (Cont.):", doc.get("notes", []))
    section("Warranties:", doc.get("warranties", []))
    section("Exclusions:", doc.get("exclusions", []))


# ----------------------------------------------------------------------------
# entry point
# ----------------------------------------------------------------------------
def build_proposal(doc, out_path):
    document = Document()
    # base style
    normal = document.styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(10)

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(2.7)      # must exceed header height so the band
    section.bottom_margin = Inches(2.3)   # lands here; footer block + logo room below
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    U.set_page_border(section)

    _build_header(section, doc["header"])
    _build_footer(section, doc["header"], total=doc.get("total"))
    render_body(document, doc)

    document.save(out_path)
    return out_path
